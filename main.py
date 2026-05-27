from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
from docxtpl import DocxTemplate
import os
import json
import zipfile
import unicodedata
import re
from datetime import datetime, timedelta
from typing import List, Optional

import stripe
from openai import OpenAI
from supabase import create_client, Client

app = FastAPI()

app.add_middleware(
    CORSMiddleware, 
    allow_origins=["*"], 
    allow_credentials=False, 
    allow_methods=["*"], 
    allow_headers=["*"]
)

# --- LLAVES MAESTRAS ---
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")

# --- LLAVES DE STRIPE (Las pondrás en Render más tarde) ---
stripe.api_key = os.environ.get("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET = os.environ.get("STRIPE_WEBHOOK_SECRET", "")

# IDs de los precios en Stripe (Ej. price_1Nxxxx...)
PRICE_ORO = os.environ.get("STRIPE_PRICE_ORO", "")
PRICE_PLATINO = os.environ.get("STRIPE_PRICE_PLATINO", "")
PRICE_BLACK = os.environ.get("STRIPE_PRICE_BLACK", "")

ia_client = OpenAI(api_key=OPENAI_API_KEY)
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# --- FUNCIONES BASE ---
def limpiar_texto(texto: str) -> str:
    if not texto: return "GENERICO"
    texto = unicodedata.normalize('NFKD', texto).encode('ASCII', 'ignore').decode('utf-8')
    return texto.upper().strip()

def limpiar_nombre_archivo(texto: str) -> str:
    if not texto: return "SN"
    return re.sub(r'[\\/*?:"<>|]', '-', str(texto)).strip()

def obtener_plantilla_y_municipio(municipio_extraido: str):
    mun = limpiar_texto(municipio_extraido)
    if "GUADALAJARA" in mun: return "Plantilla_GDL.docx", "GUADALAJARA"
    elif "ZAPOPAN" in mun: return "Plantilla_Zapopan.docx", "ZAPOPAN"
    elif "TLAQUEPAQUE" in mun: return "Plantilla_Tlaquepaque.docx", "TLAQUEPAQUE"
    elif "TLAJOMULCO" in mun: return "Plantilla_Tlajomulco.docx", "TLAJOMULCO"
    elif "TONAL" in mun: return "Plantilla_Tonala.docx", "TONALA"
    else: return "Plantilla_Generica.docx", mun if mun and mun != "GENERICO" else "GENERICO"

def extraer_texto_documento(ruta: str) -> str:
    escritura = Document(ruta)
    texto = []
    for p in escritura.paragraphs:
        if p.text.strip(): texto.append(p.text.strip())
    for table in escritura.tables:
        for row in table.rows:
            fila = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if fila: texto.append(" | ".join(fila))
    return " \n ".join(texto)

def obtener_uso_y_limite(user_id: str):
    if not user_id: return 0, 99999
    res = supabase.table("licencias").select("usos_mes, limite_mensual").eq("usada_por", user_id).execute()
    if res.data: return res.data[0].get("usos_mes", 0), res.data[0].get("limite_mensual", 3)
    return 0, 99999

def actualizar_uso(user_id: str, usos_actuales: int, nuevos_gastados: int):
    if user_id:
        supabase.table("licencias").update({"usos_mes": usos_actuales + nuevos_gastados}).eq("usada_por", user_id).execute()

PROMPT_SISTEMA = """Eres el Abogado Proyectista Jefe. Extrae los datos en JSON con la siguiente estructura estricta:
{
  "avisos": [
    {
      "escritura_numero": "", "cuenta_predial": "", "clave_catastral": "", "lugar_fecha_firma": "", "fecha_cierre": "", "fecha_resolucion": "",
      "naturaleza_acto": "", 
      "nombre_vendedor": "", "nacimiento_vendedor": "", "domicilio_vendedor": "", "generales_vendedor": "", "curp_vendedor": "", 
      "nombre_comprador": "", "nacimiento_comprador": "", "domicilio_comprador": "", "generales_comprador": "", "curp_comprador": "", 
      "ubicacion_inmueble": "", "uso_inmueble": "", "antecedentes_registro": "", 
      "valor_operacion": "", "valor_avaluo": "", "valor_catastral": "",
      "impuesto_monto": "", "total_liquidacion": "", 
      "nombre_notario": "", "notaria_numero": "", "correo_notario": "", "certificado_notario": "", 
      "clasificacion_inmueble": [], "lo_transmitido": "", "municipio_inmueble": "", "se_anexa": ""
    }
  ]
}
REGLAS DE ORO INQUEBRANTABLES (PROHIBIDO RESUMIR O INVENTAR):
1. CERTIFICADO DEL NOTARIO (CRÍTICO): En 'certificado_notario', busca al inicio el párrafo de adscripción. COPIA COMPLETO Y TEXTUAL.
2. NOMBRES CON TÍTULOS: Incluye prefijos como "Los señores", "La sociedad mercantil".
3. GENERALES EXACTAS: NO RESUMAS NADA. Copia textual el párrafo de generales.
4. RFC Y CURP: Extrae ÚNICAMENTE el código alfanumérico. OMITE el deletreo fonético.
5. VALORES MONETARIOS: Extrae Avalúo, Operación, Catastral. El 'total_liquidacion' = 'impuesto_monto'.
6. ANTECEDENTES Y UBICACIÓN: Copia linderos y antecedentes LITERAL.
7. CLASIFICACIÓN: ARREGLO ["Urbano", "Rústico", "Baldío", "Construido"].
8. SUBDIVISIONES: Un objeto en 'avisos' POR CADA INMUEBLE.
9. 'se_anexa': 'Deslinde', 'Avalúo Bancario', 'Certificado de No Propiedad', 'Certificado de no Adeudo' o 'Ninguno'."""

# --- RUTAS DE STRIPE ---
@app.post("/create-checkout-session")
async def create_checkout_session(payload: dict):
    plan = payload.get("plan")
    user_id = payload.get("user_id")
    email = payload.get("email")

    mapa_precios = {
        "Oro": PRICE_ORO,
        "Platino": PRICE_PLATINO,
        "Black": PRICE_BLACK
    }
    
    price_id = mapa_precios.get(plan)
    if not price_id:
        return {"success": False, "error": "Plan no válido o en configuración."}

    try:
        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{'price': price_id, 'quantity': 1}],
            mode='subscription',
            success_url="https://actarium.mx/terminal?pago=exito",
            cancel_url="https://actarium.mx/terminal?pago=cancelado",
            customer_email=email,
            client_reference_id=user_id, # VITAL para saber a quién le damos la licencia
        )
        return {"success": True, "url": session.url}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.post("/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        user_id = session.get("client_reference_id")
        
        # Recuperamos qué plan compró basándonos en el monto o producto
        # Simplificación: Reseteamos usos a 0 y actualizamos plan
        if user_id:
            # Determinamos el límite por plan simulado (Idealmente leemos el price_id del evento)
            limite = 50 # Default safe fallback
            nuevo_plan = "Premium"
            
            # Calculamos renovación (30 días)
            renovacion = (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d")
            
            # Actualizamos la Base de Datos
            supabase.table("licencias").update({
                "usos_mes": 0,
                "estado": "activa",
                "fecha_renovacion": renovacion
            }).eq("usada_por", user_id).execute()

    return {"status": "success"}

# --- RUTAS DE EXTRACCIÓN (Intactas y Perfeccionadas) ---
@app.post("/extraer-datos")
async def extraer_datos(file: UploadFile = File(...)):
    ruta_temporal = f"temp_{file.filename}"
    with open(ruta_temporal, "wb") as buffer: buffer.write(await file.read())
    texto_optimizado = re.sub(r'\s+', ' ', extraer_texto_documento(ruta_temporal)).strip()
    os.remove(ruta_temporal)
    
    respuesta = ia_client.chat.completions.create(
        model="gpt-4o-mini", response_format={ "type": "json_object" },
        messages=[{"role": "system", "content": PROMPT_SISTEMA}, {"role": "user", "content": texto_optimizado}]
    )
    datos_ia = json.loads(respuesta.choices[0].message.content)
    for aviso in datos_ia.get("avisos", []):
        if not aviso.get("total_liquidacion"): aviso["total_liquidacion"] = aviso.get("impuesto_monto", "")
    return datos_ia

@app.post("/generar-final")
async def generar_final(payload: dict):
    avisos = payload.get("avisos", [])
    user_id = payload.get("user_id", "")
    
    usos, limite = obtener_uso_y_limite(user_id)
    if usos >= limite: return {"success": False, "error": f"Límite de plan alcanzado ({usos}/{limite} Avisos). Adquiere un Plan."}

    archivos_generados = []
    timestamp = datetime.now().strftime("%H%M%S")
    for idx, aviso in enumerate(avisos):
        clasif = aviso.get("clasificacion_inmueble", [])
        if isinstance(clasif, str): clasif = [clasif]
        aviso["x_urbano"] = "X" if "Urbano" in clasif else " "
        aviso["x_rustico"] = "X" if "Rústico" in clasif else " "
        aviso["x_baldio"] = "X" if "Baldío" in clasif else " "
        aviso["x_construido"] = "X" if "Construido" in clasif else " "
        trans = aviso.get("lo_transmitido", "")
        aviso["x_fraccion"] = "X" if trans == "Fracción" else " "
        aviso["x_resto"] = "X" if trans == "Resto" else " "
        aviso["x_totalidad"] = "X" if trans == "Totalidad" else " "
        anexo = aviso.get("se_anexa", "Avalúo Bancario")
        aviso["x_deslinde"] = "X" if anexo == "Deslinde" else " "
        aviso["x_avaluo"] = "X" if anexo == "Avalúo Bancario" else " "
        aviso["x_certificado_no_prop"] = "X" if anexo == "Certificado de No Propiedad" else " "
        aviso["x_certificado_no_adeudo"] = "X" if anexo == "Certificado de no Adeudo" else " "

        plantilla_doc, mun_limpio = obtener_plantilla_y_municipio(aviso.get("municipio_inmueble", ""))
        if not os.path.exists(plantilla_doc): plantilla_doc = "Plantilla_Generica.docx"

        doc = DocxTemplate(plantilla_doc)
        doc.render(aviso)
        
        num_escritura = limpiar_nombre_archivo(aviso.get('escritura_numero', 'SN'))
        nombre_oficial = f"ATP_{num_escritura}_{mun_limpio}_{idx+1}.docx"
        nombre_unico = f"ATP_{num_escritura}_{mun_limpio}_{timestamp}_{idx+1}.docx" 
        ruta_local = f"temp_{nombre_unico}"
        doc.save(ruta_local)
        archivos_generados.append({"ruta_local": ruta_local, "nombre_unico": nombre_unico, "nombre_oficial": nombre_oficial, "aviso_data": aviso})

    actualizar_uso(user_id, usos, len(archivos_generados))

    if len(archivos_generados) == 1:
        archivo_data = archivos_generados[0]
        with open(archivo_data["ruta_local"], "rb") as f: supabase.storage.from_("avisos_generados").upload(archivo_data["nombre_unico"], f)
        if user_id: supabase.table("historial").insert({"user_id": user_id, "escritura": str(archivo_data["aviso_data"].get('escritura_numero', 'SN')), "acto": archivo_data["aviso_data"].get("naturaleza_acto", "-"), "vendedor": archivo_data["aviso_data"].get("nombre_vendedor", "-"), "comprador": archivo_data["aviso_data"].get("nombre_comprador", "-"), "archivo": archivo_data["nombre_unico"], "datos_json": json.dumps(archivo_data["aviso_data"])}).execute()
        os.remove(archivo_data["ruta_local"])
        return {"success": True, "archivo": archivo_data["nombre_unico"], "nombre_descarga": archivo_data["nombre_oficial"]}
    else:
        num_escritura_zip = limpiar_nombre_archivo(avisos[0].get('escritura_numero', 'SN'))
        nombre_zip_nube = f"Avisos_Subdivision_{num_escritura_zip}_{timestamp}.zip"
        with zipfile.ZipFile(nombre_zip_nube, 'w') as zipf:
            for arch in archivos_generados:
                zipf.write(arch["ruta_local"], arcname=arch["nombre_oficial"])
                with open(arch["ruta_local"], "rb") as f: supabase.storage.from_("avisos_generados").upload(arch["nombre_unico"], f)
                if user_id: supabase.table("historial").insert({"user_id": user_id, "escritura": str(arch["aviso_data"].get('escritura_numero', 'SN')), "acto": arch["aviso_data"].get("naturaleza_acto", "-"), "vendedor": arch["aviso_data"].get("nombre_vendedor", "-"), "comprador": arch["aviso_data"].get("nombre_comprador", "-"), "archivo": arch["nombre_unico"], "datos_json": json.dumps(arch["aviso_data"])}).execute()
                os.remove(arch["ruta_local"])
        with open(nombre_zip_nube, "rb") as f: supabase.storage.from_("avisos_generados").upload(nombre_zip_nube, f)
        os.remove(nombre_zip_nube)
        return {"success": True, "archivo": nombre_zip_nube, "nombre_descarga": f"Avisos_Subdivision_{num_escritura_zip}.zip"}

@app.post("/procesar-masivo")
async def procesar_masivo(archivos: List[UploadFile] = File(...), user_id: Optional[str] = Form(None), fecha_cierre: Optional[str] = Form("")):
    usos, limite = obtener_uso_y_limite(user_id)
    if usos >= limite: raise HTTPException(status_code=403, detail=f"Límite de plan alcanzado ({usos}/{limite} Avisos).")

    timestamp_lote = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_zip = f"Paquete_Avisos_Actarium_{timestamp_lote}.zip"
    total_avisos_generados = 0

    with zipfile.ZipFile(nombre_zip, 'w') as zipf:
        for archivo in archivos:
            ruta_temporal = f"temp_{archivo.filename}"
            with open(ruta_temporal, "wb") as buffer: buffer.write(await archivo.read())
            texto_optimizado = re.sub(r'\s+', ' ', extraer_texto_documento(ruta_temporal)).strip()
            os.remove(ruta_temporal)
            
            respuesta = ia_client.chat.completions.create(model="gpt-4o-mini", response_format={ "type": "json_object" }, messages=[{"role": "system", "content": PROMPT_SISTEMA}, {"role": "user", "content": texto_optimizado}])
            datos_ia = json.loads(respuesta.choices[0].message.content)
            
            for idx, aviso_ia in enumerate(datos_ia.get("avisos", [])):
                total_avisos_generados += 1
                if fecha_cierre: aviso_ia["fecha_cierre"] = fecha_cierre
                if not aviso_ia.get("total_liquidacion"): aviso_ia["total_liquidacion"] = aviso_ia.get("impuesto_monto", "")

                clasif = aviso_ia.get("clasificacion_inmueble", [])
                if isinstance(clasif, str): clasif = [clasif]
                aviso_ia["x_urbano"] = "X" if "Urbano" in clasif else " "
                aviso_ia["x_rustico"] = "X" if "Rústico" in clasif else " "
                aviso_ia["x_baldio"] = "X" if "Baldío" in clasif else " "
                aviso_ia["x_construido"] = "X" if "Construido" in clasif else " "

                trans = aviso_ia.get("lo_transmitido", "")
                aviso_ia["x_fraccion"] = "X" if trans == "Fracción" else " "
                aviso_ia["x_resto"] = "X" if trans == "Resto" else " "
                aviso_ia["x_totalidad"] = "X" if trans == "Totalidad" else " "

                anexo = aviso_ia.get("se_anexa", "Avalúo Bancario")
                aviso_ia["x_deslinde"] = "X" if anexo == "Deslinde" else " "
                aviso_ia["x_avaluo"] = "X" if anexo == "Avalúo Bancario" else " "
                aviso_ia["x_certificado_no_prop"] = "X" if anexo == "Certificado de No Propiedad" else " "
                aviso_ia["x_certificado_no_adeudo"] = "X" if anexo == "Certificado de no Adeudo" else " "
                
                mun_limpio = obtener_plantilla_y_municipio(aviso_ia.get("municipio_inmueble", ""))[1]
                plantilla_doc = obtener_plantilla_y_municipio(aviso_ia.get("municipio_inmueble", ""))[0]
                if not os.path.exists(plantilla_doc): plantilla_doc = "Plantilla_Generica.docx"

                doc = DocxTemplate(plantilla_doc)
                doc.render(aviso_ia)
                
                num_escritura = limpiar_nombre_archivo(aviso_ia.get('escritura_numero', 'SN'))
                nombre_limpio = limpiar_nombre_archivo(archivo.filename.replace(".docx", ""))
                nombre_archivo_final = f"ATP_{num_escritura}_{mun_limpio}_{idx+1}.docx" if len(datos_ia.get("avisos", [])) > 1 else f"ATP_{num_escritura}_{mun_limpio}.docx"
                
                nombre_unico_nube = f"ATP_{num_escritura}_{mun_limpio}_{datetime.now().strftime('%H%M%S')}_{idx+1}.docx"
                ruta_local = f"temp_{nombre_unico_nube}"
                doc.save(ruta_local)
                
                with open(ruta_local, "rb") as f: supabase.storage.from_("avisos_generados").upload(nombre_unico_nube, f)
                if user_id: supabase.table("historial").insert({"user_id": user_id, "escritura": str(num_escritura), "acto": aviso_ia.get("naturaleza_acto", "-"), "vendedor": aviso_ia.get("nombre_vendedor", "-"), "comprador": aviso_ia.get("nombre_comprador", "-"), "archivo": nombre_unico_nube, "datos_json": json.dumps(aviso_ia)}).execute()
                zipf.write(ruta_local, arcname=f"{mun_limpio}/{nombre_limpio}/{nombre_archivo_final}")
                os.remove(ruta_local)

    if user_id and total_avisos_generados > 0: actualizar_uso(user_id, usos, total_avisos_generados)
    return FileResponse(nombre_zip, filename=nombre_zip, media_type="application/zip")